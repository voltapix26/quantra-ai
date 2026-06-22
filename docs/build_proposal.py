# -*- coding: utf-8 -*-
"""Builds a business proposal + cover letter for Quantra AI (by Eshan Thanvi)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x1B, 0x9E, 0x6B)
INK = RGBColor(0x18, 0x1D, 0x26)
MUTE = RGBColor(0x55, 0x5F, 0x6B)
BLUE = RGBColor(0x2F, 0x5C, 0xC8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
AUTHOR = "Eshan Thanvi"
EMAIL = "eshanthanvi@gmail.com"
URL = "quantra-ai.onrender.com"


def new_doc():
    d = Document()
    s = d.styles['Normal']; s.font.name = 'Segoe UI'; s.font.size = Pt(10.5); s.font.color.rgb = INK
    sec = d.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.8); sec.left_margin = sec.right_margin = Inches(0.9)
    return d


def para(d, text='', size=10.5, color=INK, bold=False, italic=False, after=6, before=0, align=None, line=1.25):
    p = d.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(after); pf.space_before = Pt(before); pf.line_spacing = line
    if align is not None: p.alignment = align
    for seg in (text if isinstance(text, list) else [(text, bold)]):
        t, b = seg if isinstance(seg, tuple) else (seg, bold)
        r = p.add_run(t); r.font.size = Pt(size); r.font.color.rgb = color; r.font.bold = b; r.font.italic = italic; r.font.name = 'Segoe UI'
    return p


def band(d):
    b = d.add_paragraph(); b.paragraph_format.space_after = Pt(6)
    pb = b._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '8'); bt.set(qn('w:color'), '1B9E6B'); bt.set(qn('w:space'), '1')
    bd.append(bt); pb.append(bd)


def h(d, text):
    p = d.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = GREEN; r.font.name = 'Segoe UI'


def bullet(d, label, body=''):
    p = d.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label); r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = INK
    if body:
        r2 = p.add_run(' — ' + body); r2.font.size = Pt(10.5); r2.font.color.rgb = MUTE


def letterhead(d, doctype):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    r = p.add_run('Quantra'); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = INK; r.font.name = 'Segoe UI'
    r2 = p.add_run('AI'); r2.font.size = Pt(22); r2.font.bold = True; r2.font.color.rgb = GREEN; r2.font.name = 'Segoe UI'
    para(d, doctype, 11.5, GREEN, bold=True, after=1)
    para(d, [(f"{AUTHOR}  ·  ", False), (EMAIL, False), ("  ·  ", False), (URL, False)], 9, MUTE, after=4)
    band(d)


# ============================================================ COVER LETTER
d = new_doc()
letterhead(d, "Cover Letter")
para(d, "22 June 2026", 10.5, MUTE, after=10)
para(d, "Dear Review Committee,", 10.5, INK, after=8)

para(d, [("I am pleased to present ", False), ("Quantra AI", True),
         (", a professional-grade, multi-tenant market-analysis platform I designed and built end to end. "
          "Quantra brings the kind of capability normally locked behind $20,000-a-year terminals — real-time "
          "data, advanced charting, calibrated forecasting and an AI analyst — into a single, accessible web "
          "application that anyone can use from a browser or phone.", False)], 10.5, INK, after=8)

para(d, [("The platform is ", False), ("live in production", True),
         (f" at {URL}, deployed with continuous delivery. It already spans 25 world stock exchanges, 32 global "
          "indices, ETFs, crypto, FX and commodities; a 16-study technical suite; an AI verdict engine; and a "
          "probabilistic forecast model independently calibrated to ~80% band coverage on live data — a level of "
          "statistical honesty rarely offered to retail users.", False)], 10.5, INK, after=8)

para(d, [("Beyond analytics, Quantra now includes portfolio tracking, paper trading, a community layer, and a "
          "compliant ", False), ("bring-your-own-broker", True),
         (" trading module in which users link their own regulated brokerage — Quantra routes orders but never "
          "custodies funds. Every design decision balances capability with safety, transparency and regulatory "
          "awareness.", False)], 10.5, INK, after=8)

para(d, [("The attached proposal details the problem, product, technology, market and roadmap. I would welcome "
          "the opportunity to discuss how Quantra AI can be taken to its next stage of growth.", False)],
     10.5, INK, after=12)

para(d, "Sincerely,", 10.5, INK, after=2)
para(d, AUTHOR, 11.5, INK, bold=True, after=0)
para(d, "Founder & Builder, Quantra AI", 9.5, MUTE, after=0)
d.save(r"C:\Users\eshan\quantra-terminal\docs\Quantra_AI_Cover_Letter.docx")
print("OK cover letter")

# ============================================================ PROPOSAL
d = new_doc()
letterhead(d, "Business Proposal")
para(d, [("Prepared by ", False), (AUTHOR, True), ("  ·  22 June 2026", False)], 9.5, MUTE, after=8)

h(d, "1 · Executive summary")
para(d, [("Quantra AI", True),
         (" is a live, multi-tenant SaaS platform that democratizes professional market analysis. It delivers "
          "real-time and global market data, an advanced charting and indicator suite, an AI analyst, and "
          "probabilistic forecasts calibrated to ~80% accuracy — accessible from any browser or phone, with no "
          "expensive terminal required.", False)], after=6)

h(d, "2 · The problem")
bullet(d, "Cost", "professional tools (Bloomberg, Refinitiv) cost $20k+/year and are out of reach for retail users, students and small funds.")
bullet(d, "Fragmentation", "charts, news, portfolio and forecasts live in separate, disconnected tools.")
bullet(d, "Trust", "retail 'projections' are typically decorative and uncalibrated; data quality is opaque.")

h(d, "3 · The solution")
para(d, "One terminal that unifies global coverage, real-time data, professional charting, calibrated forecasting "
        "and an AI analyst — with honesty and transparency built in.", after=6)
bullet(d, "Global coverage", "25 exchanges, 32 indices, ETFs, crypto, FX and commodities, in native currency.")
bullet(d, "Pro charting", "16 studies (Ichimoku, Supertrend, VWAP, Bollinger, ADX and more) on line & candlestick charts.")
bullet(d, "Calibrated forecasts", "Monte-Carlo bands proven to ~80% coverage, shown live on a public Track Record page.")
bullet(d, "AI analyst & brief", "grounded per-asset analysis and a personalized daily digest.")
bullet(d, "Full workflow", "portfolio, paper trading, alerts, calendar and community — plus broker-linked trading.")

h(d, "4 · Technology")
para(d, "A lean, resilient Node.js platform acting as market-data proxy and application server, with optional "
        "integrations that activate only when configured. Multi-tenant accounts, super-admin oversight, a PWA with "
        "push notifications, and continuous deployment. Paid data sources are key-gated with graceful fallback, so "
        "the product degrades gracefully rather than breaking.", after=6)

h(d, "5 · Market & differentiation")
para(d, "Quantra targets retail investors, students, analysts and small funds priced out of incumbent terminals. "
        "Its differentiators are calibrated (provable) forecasts, radical data transparency, breadth of global "
        "coverage, and an AI analyst — delivered at a fraction of incumbent cost.", after=6)

h(d, "6 · Business model")
bullet(d, "Subscription tiers", "Free, Pro and Ultimate, billed via Stripe (integrated).")
bullet(d, "Compliant trading", "bring-your-own-broker module; users custody funds at their own regulated broker.")
bullet(d, "Data partners", "aggregator feeds (Twelve Data, Finnhub, Polygon) scale coverage without custodial risk.")

h(d, "7 · Traction & status")
bullet(d, "Live in production", "deployed with GitHub continuous delivery; current build v46.")
bullet(d, "Feature-complete core", "the full 10-feature roadmap plus global markets, indicators and broker trading shipped.")
bullet(d, "Proven calibration", "~80% realised band coverage measured on live forward outcomes.")

h(d, "8 · Roadmap")
bullet(d, "Data", "real-time global feeds via aggregator keys (Twelve Data / Polygon).")
bullet(d, "Brokers", "additional broker adapters (e.g. Zerodha/Upstox for India) under compliance review.")
bullet(d, "Mobile & growth", "native app packaging, onboarding and conversion optimisation.")

h(d, "9 · The ask")
para(d, "I am seeking partnership, mentorship or investment to scale Quantra AI — funding real-time data licensing, "
        "broker integrations and growth. I welcome the opportunity to present a live demonstration.", after=8)

para(d, [("Contact:  ", True), (f"{AUTHOR}  ·  {EMAIL}  ·  {URL}", False)], 10, MUTE, after=0)
d.save(r"C:\Users\eshan\quantra-terminal\docs\Quantra_AI_Proposal.docx")
print("OK proposal")
