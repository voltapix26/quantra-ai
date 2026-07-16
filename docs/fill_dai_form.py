# Fill the Dubai Angel Investors application form with Quantra AI's REAL data.
# Honest by design: screening Q3 (traction) and Q4 (team size) are answered NO,
# because they are false today and DAI verifies both.
import os, docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

LOGO = r'C:\Users\eshan\AppData\Local\Temp\claude\C--Program-Files-Git\6ccb22d2-c064-49e1-b626-f421b9250a09\scratchpad\quantra-logo.png'

SRC = r'C:\Users\eshan\Downloads\2017%20Sep%20DAI%20Application%20Form.docx'
OUT_DIR = os.path.join(os.path.expanduser('~'), 'Desktop', 'Quantra AI')
os.makedirs(OUT_DIR, exist_ok=True)
OUT = os.path.join(OUT_DIR, 'Quantra_AI_DAI_Application.docx')

d = docx.Document(SRC)
GREEN = RGBColor(0x0E, 0x9F, 0x6E); RED = RGBColor(0xC0, 0x27, 0x18)

def answer(p, ans, note=''):
    """Replace the 'YES / NO' in a screening question with a bolded answer."""
    full = p.text
    if 'YES / NO' not in full:
        return False
    q = full.split('YES / NO')[0].rstrip()
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r1 = p.add_run(q + '   ')
    r1.font.size = Pt(10)
    r2 = p.add_run(ans)
    r2.bold = True; r2.font.size = Pt(11)
    r2.font.color.rgb = GREEN if ans == 'YES' else RED
    if note:
        r3 = p.add_run('  — ' + note)
        r3.italic = True; r3.font.size = Pt(8.5)
        r3.font.color.rgb = RGBColor(0x55, 0x60, 0x72)
    return True

# ---- screening answers (honest) ----
ANS = {
    'client or consumer facing': ('YES', 'Consumer-facing web + PWA + Android markets terminal.'),
    'minimum viable product': ('YES', 'Full product LIVE in production: quantra-ai.onrender.com'),
    'initial traction': ('NO', 'Pre-launch: product complete + public accuracy ledger running, but no user acquisition started yet.'),
    'more than one full time team member': ('NO', 'Solo founder today; first hires are the primary use of this round.'),
    'active in the MENA region': ('YES', 'Dubai-based founder; DFM, Tadawul + Gulf coverage live; Arabic/RTL shipped.'),
    'already raised some money': ('NO', 'Bootstrapped — no outside capital to date.'),
    'have a cofounder': ('NO', 'Actively seeking a co-founder / first full-time engineer.'),
}
for p in d.paragraphs:
    for key, (a, note) in ANS.items():
        if key.lower() in p.text.lower() and 'YES / NO' in p.text:
            answer(p, a, note); break

# ---- logo + company name header ----
for p in d.paragraphs:
    if 'ADD LOGO' in p.text:
        for r in list(p.runs): r._element.getparent().remove(r._element)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(LOGO):
            p.add_run().add_picture(LOGO, width=Inches(3.2))
for p in d.paragraphs:
    if 'INSERT NAME OF COMPANY' in p.text:
        for r in list(p.runs): r._element.getparent().remove(r._element)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('QUANTRA AI'); r.bold = True; r.font.size = Pt(24)

def setcell(tbl, row, val):
    c = tbl.rows[row].cells[1]
    c.text = ''
    r = c.paragraphs[0].add_run(val); r.font.size = Pt(10)

# ---- table 0: company details ----
t = d.tables[0]
setcell(t, 1, 'Quantra AI')
setcell(t, 2, 'https://quantra-ai.onrender.com   ·   Public accuracy ledger: /track-record.html')
setcell(t, 3, 'Android: APK built & installable (Play Store listing pending). iOS/Android: installable PWA live. No iOS App Store build yet.')
setcell(t, 4, 'June 2026')
setcell(t, 5, '[YOUR REGISTERED ADDRESS — please complete]')
setcell(t, 6, 'United Arab Emirates / Dubai   [confirm — entity not yet incorporated]')
setcell(t, 7, 'Same as above')
setcell(t, 8, 'Same as above (single founder, remote-first)')

# ---- table 1: fundraising ----
t = d.tables[1]
setcell(t, 1, 'US$ 2,500,000 (seed)')
setcell(t, 2, 'US$ 2,000,000')
setcell(t, 3, 'US$ 3,000,000')
setcell(t, 4, 'Not set — open to discussion with the lead investor')
setcell(t, 5, 'US$ 0 — no commitments yet')
setcell(t, 6, 'US$ 0 — fully bootstrapped to date')
setcell(t, 7, 'Not set')

# ---- elevator pitch (<=200 chars) ----
ELEV = ('Quantra AI: the honest markets terminal. Live AI analysis across 8 asset classes where every '
        'projection is publicly graded against reality on a tamper-evident ledger. Probabilities, never promises.')
assert len(ELEV) <= 200, f'elevator too long: {len(ELEV)}'

DETAIL = (
"PROBLEM. A pro markets terminal costs US$25,000+/yr, so the fast-growing retail investor class across the GCC and "
"South Asia runs on fragmented apps and Telegram tipsters — an industry selling certainty that cannot exist: "
"guaranteed calls, back-filled win rates, cherry-picked screenshots. Nobody publishes an auditable record of when "
"they were wrong, and regional regulators are tightening on exactly this.\n\n"
"PRODUCT (live). One terminal: crypto (tick-by-tick), stocks on 24 exchanges incl. DFM, Tadawul, NSE, ETFs, "
"commodities, CME futures, indices, FX, US options, Web3 dashboard — 24 currencies, Arabic/RTL. Plus AI analysis in "
"plain language, a probabilistic Quantra Score, calibrated projection bands, a movers radar (odds of 2-40% moves, "
"1h-30d) with alerts, paper trading, portfolio risk, a developer API. Web + PWA + Android. Free/Pro/Ultimate plans "
"and Stripe billing built, one switch from live.\n\n"
"MOAT — verifiable honesty. Every projection is a probability with the downside shown, each publicly graded against "
"the real outcome on a SHA-256 hash-chained ledger anyone can audit: today 29 straight days, 19,550 graded "
"projections, no back-filling. Bands self-calibrate to measured outcomes. Rivals copy features in months; they cannot "
"copy months of public, verified accuracy.\n\n"
"TRACTION — plainly. Product complete and in production (CI every release, self-diagnostics every 12 min), but user "
"acquisition has NOT started: no active users, no revenue. Our verifiable asset is the product plus the accuracy "
"record compounding daily. We answer NO to Q3 and Q4 rather than overstate.\n\n"
"TEAM. Solo founder (Eshan Thanvi), built AI-assisted — a Series-A feature surface at seed cost. First hires and a "
"co-founder search are the primary use of this round.\n\n"
"ASK. US$2.5M seed: 35% team; 30% licensed exchange data (Gulf + Indian F&O adapters written, awaiting paid feeds) "
"+ infra; 20% GCC/India go-to-market; 10% ADGM regulatory footing; 5% reserve.\n\n"
"Demo + public proof: quantra-ai.onrender.com"
)
assert len(DETAIL) <= 2000, f'detailed too long: {len(DETAIL)}'

def fillbox(tbl, text, size=9):
    c = tbl.rows[0].cells[0]; c.text = ''
    first = True
    for para in text.split('\n\n'):
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        r = p.add_run(para); r.font.size = Pt(size)

fillbox(d.tables[2], ELEV, 11)
fillbox(d.tables[3], DETAIL, 8.5)

# ---- table 4: documents ----
DOCS = {
    'Completed Application': ('Yes', ''),
    'Pitch Deck': ('Yes', ''),
    'Business Plan': ('No', 'Yes'),
    'Financial Model': ('No', 'Yes'),
    'Due Diligence': ('No', 'Yes'),
    'Product Video': ('No', 'Yes'),
    'Customer Testimonials': ('No', 'No'),
}
t = d.tables[4]
for row in t.rows:
    label = row.cells[0].text.strip()
    for key, (sub, able) in DOCS.items():
        if key.lower() in label.lower():
            row.cells[1].text = sub
            if able: row.cells[2].text = able
            break
    if 'Other' in label:
        row.cells[1].text = ('Live product (quantra-ai.onrender.com) · Public tamper-evident accuracy ledger '
                             '(/track-record.html) · Public system status (/status.html) · Technical documentation '
                             '(PRD, TRD, architecture, schema, ops + billing runbooks) · Android APK')

d.save(OUT)
print('elevator chars:', len(ELEV), '| detailed chars:', len(DETAIL))
print('SAVED:', OUT)
