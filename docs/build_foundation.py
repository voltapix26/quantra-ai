# Builds "Quantra_AI_Foundation_Documents.docx" (PRD/TRD/App Flow/UIUX/Schema/Plan)
# from docs/foundation/*.md — light markdown parsing: headings, tables, lists, code.
import os, re, glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "foundation")
OUT_DIR = os.path.join(os.path.expanduser("~"), "Desktop", "Quantra AI")
os.makedirs(OUT_DIR, exist_ok=True)
OUT = os.path.join(OUT_DIR, "Quantra_AI_Foundation_Documents.docx")

MINT = RGBColor(0x0E, 0x9F, 0x6E)
DARK = RGBColor(0x0A, 0x0F, 0x1C)
MUT = RGBColor(0x55, 0x60, 0x72)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5)

# cover
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\nQuantra AI"); r.font.size = Pt(40); r.bold = True; r.font.color.rgb = DARK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Foundation Documents\nPRD · TRD · App Flow · UI/UX Brief · Backend Schema · Implementation Plan")
r.font.size = Pt(16); r.font.color.rgb = MINT
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Status-audited against production (quantra-ai.onrender.com) · July 2026\nPrepared by Eshan Thanvi")
r.font.size = Pt(11); r.font.color.rgb = MUT
doc.add_page_break()

def add_table(rows):
    if not rows: return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j >= len(t.rows[i].cells): continue
            c = t.rows[i].cells[j]
            c.text = ""
            r = c.paragraphs[0].add_run(cell)
            r.font.size = Pt(9.5)
            if i == 0: r.bold = True
    doc.add_paragraph()

for path in sorted(glob.glob(os.path.join(SRC, "*.md"))):
    lines = open(path, encoding="utf-8").read().splitlines()
    tbl, in_code = [], False
    def flush():
        global tbl
        if tbl: add_table(tbl); tbl = []
    for ln in lines:
        if ln.strip().startswith("```"):
            flush(); in_code = not in_code; continue
        if in_code:
            p = doc.add_paragraph(); r = p.add_run(ln)
            r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = MUT
            p.paragraph_format.space_after = Pt(0)
            continue
        if ln.startswith("|"):
            cells = [c.strip().replace("**", "") for c in ln.strip().strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells): continue
            tbl.append(cells); continue
        flush()
        if ln.startswith("# "):
            doc.add_page_break() if doc.paragraphs[-1].text else None
            h = doc.add_heading(ln[2:].strip(), level=1)
            for r in h.runs: r.font.color.rgb = DARK
        elif ln.startswith("## "):
            h = doc.add_heading(ln[3:].strip(), level=2)
            for r in h.runs: r.font.color.rgb = MINT
        elif ln.startswith("### "):
            doc.add_heading(ln[4:].strip(), level=3)
        elif ln.startswith("> "):
            p = doc.add_paragraph(); r = p.add_run(ln[2:].strip()); r.italic = True; r.font.color.rgb = MUT
        elif ln.strip().startswith(("- ", "* ")):
            txt = ln.strip()[2:].replace("**", "")
            doc.add_paragraph(txt, style="List Bullet")
        elif re.match(r"^\d+\.\s", ln.strip()):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", ln.strip()).replace("**", ""), style="List Number")
        elif ln.strip():
            doc.add_paragraph(ln.replace("**", ""))
    flush()

doc.save(OUT)
print("SAVED:", OUT)
