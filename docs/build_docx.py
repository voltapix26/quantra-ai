# -*- coding: utf-8 -*-
"""Builds docs/Quantra_AI_Version_History.docx from the changelog content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x1B, 0x9E, 0x6B)
INK = RGBColor(0x1A, 0x1F, 0x28)
MUTE = RGBColor(0x5C, 0x66, 0x72)
BLUE = RGBColor(0x2F, 0x5C, 0xC8)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Segoe UI'; st.font.size = Pt(10.5); st.font.color.rgb = INK
for m in ('top', 'bottom'):
    setattr(doc.sections[0], m + '_margin', Inches(0.8))
doc.sections[0].left_margin = doc.sections[0].right_margin = Inches(0.9)


def shade(p, color="1B9E6B"):
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), color)
    pPr.append(sh)


def para(text='', size=10.5, color=INK, bold=False, italic=False, after=4, before=0, align=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color
    r.font.bold = bold; r.font.italic = italic; r.font.name = 'Segoe UI'
    return p, r


def phase(title):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)
    shade(p, "0B0E14")
    r = p.add_run('  ' + title); r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.name = 'Segoe UI'


def release(ver, name, meta):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(ver + '  '); r.font.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = GREEN
    r2 = p.add_run('— ' + name); r2.font.bold = True; r2.font.size = Pt(11.5); r2.font.color.rgb = INK
    mp, mr = para(meta, 8.5, MUTE, italic=True, after=3)


def bullet(head, sub=''):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(head); r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = INK
    if sub:
        r2 = p.add_run('  ' + sub); r2.font.size = Pt(9.5); r2.font.color.rgb = MUTE


# ---- Title block ----
_, r = para('Quantra AI', 30, INK, bold=True, after=0)
para('Version History & Changelog', 15, GREEN, bold=True, after=8)
para('Multi-tenant SaaS market-analysis terminal  ·  Phases 1–6  ·  v1.0.0 → v5.7.0',
     10.5, MUTE, after=2)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
for t, c, b in [('Live: ', MUTE, False), ('quantra-ai.onrender.com', BLUE, True),
                ('     Repo: ', MUTE, False), ('github.com/voltapix26/quantra-ai', BLUE, True)]:
    rr = p.add_run(t); rr.font.size = Pt(9.5); rr.font.color.rgb = c; rr.font.bold = b; rr.font.name = 'Segoe UI'
para('Current build: Assets v39 · Service Worker quantra-v21 · 2026-06-19', 9.5, MUTE, after=4)
b = doc.add_paragraph(); b.paragraph_format.space_after = Pt(0)
pb = b._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '8'); bt.set(qn('w:color'), '1B9E6B'); bt.set(qn('w:space'), '1')
bd.append(bt); pb.append(bd)

DATA = [
 ("Phase 6 — Market Status & Real-Time Feeds", [
   ("v5.7.0", "Faster, honest live feed", "2026-06-19 · 32f283c · current", [
     ("Real-time detail poll 4s → 2.5s; board refresh 15s → 10s.", ""),
     ("Accurate freshness label:", "“live” (open + fresh), “delayed feed” (open + stale), “at last close” (market closed).")]),
   ("v5.6.0", "Per-exchange holiday calendar", "2026-06-19 · 49a75df", [
     ("Authoritative US status via Finnhub market-status", "— flags closures such as Juneteenth → “Closed · Juneteenth”."),
     ("Curated 2026 holiday calendar for 13 major non-US exchanges", "(UK, DE, FR, IT, ES, CH, JP, HK, AU, CA, SG, India NSE/BSE)."),
     ("Threaded through board, detail badge, Discover and Portfolio.", "")]),
   ("v5.5.0", "Stale-session fix (timezone-aware)", "2026-06-19 · c27b08a", [
     ("Open/closed computed from session time-of-day in the exchange timezone", "— immune to Yahoo’s day-stale window.")]),
   ("v5.4.0", "Real-time US feed + status dots everywhere", "2026-06-19 · 90df4dc", [
     ("Finnhub real-time US quotes override Yahoo with graceful fallback;", "session window returned by /api/price."),
     ("Open/closed dots added to Discover movers and Portfolio holdings.", "")]),
   ("v5.3.0", "Open/close countdown", "2026-06-19 · c1bad4a", [
     ("Badge shows “closes in 3h 5m” / “opens in 14h”, live every 20s, across weekends.", "")]),
   ("v5.2.0", "Live-flipping status", "2026-06-19 · 299ddf3", [
     ("Dots and badge re-evaluate every 20s, flipping exactly at the open/close boundary.", "")]),
   ("v5.1.0", "Per-row board status dots", "2026-06-19 · 98d5560", [
     ("Green = open, grey = closed per exchange; crypto always-on, FX weekday-aware.", "")]),
   ("v5.0.0", "Market open/closed badge", "2026-06-19 · b1c1599", [
     ("Per-exchange session badge: Market open / Pre-market / After-hours / Closed.", "")]),
 ]),
 ("Phase 5 — Technical Indicator Suite", [
   ("v4.2.0", "Enlarge mode + dynamic legend", "2026-06-19 · 2b49273", [
     ("Full-width / taller chart toggle (proportional viewBox) + legend of active overlays.", "")]),
   ("v4.1.0", "Full indicator suite", "2026-06-19 · 215728b", [
     ("Overlays:", "Ichimoku Cloud, Supertrend, VWAP, Keltner & Donchian channels."),
     ("Panes:", "ADX, Williams %R, CCI."),
     ("Shared engine renders every study on both line and candlestick charts; saved in layouts.", "")]),
   ("v4.0.0", "Adaptive averages + Stochastic", "2026-06-19 · 43e075a", [
     ("McGinley Dynamic + EMA(21) + Parabolic SAR overlays and a Stochastic (14·3) pane.", "")]),
 ]),
 ("Phase 4 — Global Markets & Accurate Pricing", [
   ("v3.10.0", "Twelve Data + absolute change", "2026-06-19 · 0ef471d", [
     ("Twelve Data (gated) for fresher global/Gulf quotes; absolute up/down amount beside the %.", "")]),
   ("v3.9.0", "Data freshness transparency", "2026-06-19 · e345a54", [
     ("“Price as of <time> <tz>” line so delayed free feeds are visible, not wrong.", "")]),
   ("v3.8.0", "Regional ETFs + Discover market filter", "2026-06-19 · 935d33b", [
     ("Country/region ETFs + per-exchange breadth/movers/heatmap on Discover.", "")]),
   ("v3.7.0", "32 global indices", "2026-06-19 · 789ca78", [
     ("TSX, Bovespa, KOSPI, TWSE, STI, SMI, IBEX, FTSE MIB, AEX, SSE, Bank Nifty, TASI, JSE …", "shown as raw points.")]),
   ("v3.6.0", "25 world stock exchanges", "2026-06-19 · 878e4ca", [
     ("US → South Africa; currency follows the exchange (16 ccys) with pence/cents handling.", "")]),
   ("v3.5.0", "Stock market bifurcation", "2026-06-19 · 2815a77", [
     ("Exchange selector (US / NSE / BSE / Europe / UAE / HK); currency auto-follows.", "")]),
   ("v3.4.0", "Speed: keep-warm + faster board", "2026-06-19 · aaebd38", [
     ("Keep-warm self-ping removes free-tier cold start (~50s → instant); board 25s → 15s.", "")]),
   ("v3.3.0", "Fix board 24h % (MSFT bug)", "2026-06-19 · 4a301fa", [
     ("chartPreviousClose (~1 month old) → now prior-session close, session-aware.", "")]),
 ]),
 ("Phase 3 — Premium Data & Forecast Accuracy", [
   ("v3.2.0", "Live calibration widget", "2026-06-19 · c8f3119", [
     ("Server measures realised 80%-band coverage from real forward outcomes (provable).", "")]),
   ("v3.1.0", "Forecast calibration to 80%", "2026-06-19 · 9da53c6", [
     ("Bootstrap MC + drift shrinkage + MC-median path → 80.1% realised vs 74.4% before.", "")]),
   ("v3.0.0", "Premium data feeds", "2026-06-19 · 970189b", [
     ("FMP economic calendar + marketaux news (source badges, sentiment); key-gated.", "")]),
 ]),
 ("Phase 2 — 10-Feature SaaS Roadmap", [
   ("v2.10.0", "Community", "2026-06-19 · a943691", [("Shared trade ideas (post/upvote/delete) + paper-trading leaderboard.", "")]),
   ("v2.9.0", "Paper trading", "2026-06-19 · 80e1a56", [("Simulated cash, live buy/sell, P&L, trade history, research journal.", "")]),
   ("v2.8.0", "Bollinger + bar replay + saved layouts", "2026-06-19 · a662f39", [("Replay (step/play/scrub) + account-synced layouts.", "")]),
   ("v2.7.0", "AI daily brief", "2026-06-19 · 81f6056", [("Personalized digest: movers, breadth, upcoming earnings.", "")]),
   ("v2.6.0", "PWA + push", "2026-06-19 · 4eb20e2", [("Installable app + offline shell + VAPID web-push alerts.", "")]),
   ("v2.5.0", "Market calendar", "2026-06-19 · 9854558", [("Earnings, IPOs & economic events (Finnhub + graceful fallback).", "")]),
   ("v2.4.0", "Market discovery", "2026-06-19 · 66c012d", [("Heatmap + movers + breadth gauge across all asset classes.", "")]),
   ("v2.3.0", "Monitored alerts + email", "2026-06-19 · 8bcd130", [("Server-side engine fires even when the tab is closed.", "")]),
   ("v2.2.0", "Ask Quantra AI analyst", "2026-06-19 · 56996df", [("Grounded chat per asset (LLM + local fallback), rate-limited.", "")]),
   ("v2.1.0", "Portfolio tracker", "2026-06-19 · 3fc73dd", [("Holdings, live P&L, allocation; universal /api/price endpoint.", "")]),
 ]),
 ("Phase 1 — Foundation & Live Data (2026-06-18)", [
   ("v1.8.0", "Admin oversight & analytics", "089e84c · b6129de · fc9f8e9 · ecdd330", [
     ("Audit log + admin panel (metadata only, never passwords); footfall analytics.", "")]),
   ("v1.7.0", "Alerts & drawing tools", "4ab96d9 · af72a91", [("Price alerts + trendline / level / Fibonacci tools.", "")]),
   ("v1.6.0", "Indicator panes & chart types", "d46ccf1", [("RSI / MACD / Volume panes, patterns, Heikin Ashi + Area.", "")]),
   ("v1.5.0", "Tick-by-tick stocks", "621e9c6 · 4b76a12 · a6dfc9b", [("Live seconds chart (Finnhub ~1s) + WS → SSE relay.", "")]),
   ("v1.4.0", "Cloud-resilient live crypto", "c9ccd3f · 8581ed5", [("Coinbase WS (Binance geoblocks Render); sub-second ticks.", "")]),
   ("v1.3.0", "Candles + dated projections", "9c4ff12", [("Candlestick toggle + dated, checkable Monte-Carlo projections.", "")]),
   ("v1.2.0", "Live data backbone", "5971d34", [("Crypto WS + Finnhub real-time quotes & news.", "")]),
   ("v1.1.0", "Responsive mobile", "812f217 · fe14f60 · 4ebc1f1", [("Phone breakpoints; screener + notes layouts.", "")]),
   ("v1.0.0", "Initial deploy", "147945d · af2c447 · a60ba33", [("First public deploy + share launcher + crash guards.", "")]),
 ]),
]

for ptitle, rels in DATA:
    phase(ptitle)
    for ver, name, meta, bs in rels:
        release(ver, name, meta)
        for h, sub in bs:
            bullet(h, sub)

phase("Reference")
para("Data sources", 11, GREEN, bold=True, before=4, after=2)
para("CoinGecko / CoinPaprika (crypto) · Coinbase (OHLC + WS) · Yahoo Finance (stocks/ETF/index/"
     "commodity/FX) · Finnhub (real-time US, news, calendars, market-status, WS) · Twelve Data (gated "
     "global) · FMP (gated economic calendar) · marketaux (gated news) · open.er-api (FX). All paid "
     "sources are key-gated with graceful fallback.", 9.5, MUTE, after=6)
para("Optional environment keys", 11, GREEN, bold=True, after=2)
para("ANTHROPIC_API_KEY · RESEND_API_KEY · VAPID_PUBLIC/PRIVATE_KEY + VAPID_SUBJECT · FINNHUB_API_KEY "
     "· TWELVEDATA_API_KEY · FMP_API_KEY · MARKETAUX_API_KEY · COINGECKO_KEY · SUPER_ADMINS.", 9.5, MUTE)

doc.save(r"C:\Users\eshan\quantra-terminal\docs\Quantra_AI_Version_History.docx")
print("OK docx")
