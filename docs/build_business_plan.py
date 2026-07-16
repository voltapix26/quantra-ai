# Quantra AI — Business Plan (docx). Honest by construction: no invented traction,
# projections labelled as projections, gaps stated plainly.
import os, docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.expanduser('~'), 'Desktop', 'Quantra AI')
os.makedirs(OUT_DIR, exist_ok=True)
OUT = os.path.join(OUT_DIR, 'Quantra_AI_Business_Plan.docx')
LOGO = r'C:\Users\eshan\AppData\Local\Temp\claude\C--Program-Files-Git\6ccb22d2-c064-49e1-b626-f421b9250a09\scratchpad\quantra-logo.png'

MINT = RGBColor(0x0E, 0x9F, 0x6E); DARK = RGBColor(0x0A, 0x0F, 0x1C); MUT = RGBColor(0x55, 0x60, 0x72)
RED = RGBColor(0xB0, 0x00, 0x20)
d = docx.Document()
st = d.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)

def H1(t):
    h = d.add_heading(t, level=1)
    for r in h.runs: r.font.color.rgb = DARK
def H2(t):
    h = d.add_heading(t, level=2)
    for r in h.runs: r.font.color.rgb = MINT
def P(t, size=10.5, italic=False, color=None, bold=False):
    p = d.add_paragraph(); r = p.add_run(t)
    r.font.size = Pt(size); r.italic = italic; r.bold = bold
    if color: r.font.color.rgb = color
    return p
def B(t): d.add_paragraph(t, style='List Bullet')
def table(rows, widths=None):
    t = d.add_table(rows=len(rows), cols=len(rows[0])); t.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]; c.text = ''
            r = c.paragraphs[0].add_run(str(val)); r.font.size = Pt(9.5)
            if i == 0: r.bold = True
    d.add_paragraph()

# ---- cover ----
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO): p.add_run().add_picture(LOGO, width=Inches(3.4))
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\nBusiness Plan'); r.font.size = Pt(30); r.bold = True; r.font.color.rgb = DARK
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('The honest markets terminal — live AI analysis, publicly graded against reality')
r.font.size = Pt(13); r.font.color.rgb = MINT
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Seed round: US$2.5M  ·  July 2026\nEshan Thanvi, Founder  ·  quantra.aio@proton.me  ·  quantra-ai.onrender.com')
r.font.size = Pt(10.5); r.font.color.rgb = MUT
d.add_page_break()

# ---- 1 exec summary ----
H1('1. Executive Summary')
P('Quantra AI is a live, multi-asset market-analysis terminal for retail investors across the Gulf and South Asia. '
  'It gives an individual investor the capability set of a US$25,000/year institutional terminal — live data across eight '
  'asset classes, AI-generated analysis in plain language, and probabilistic price projections — with one structural '
  'difference: every projection Quantra makes is publicly graded against what actually happened, on a tamper-evident '
  'ledger anyone can audit.')
P('The product is complete and in production today. What it does not yet have is users or revenue: acquisition has not '
  'started. We are raising a US$2.5M seed to buy licensed exchange data, put the first team in place, establish UAE '
  'regulatory footing, and begin go-to-market in the GCC and India.', bold=True)
H2('Status at a glance')
table([
    ['Item', 'Position today'],
    ['Product', 'Live in production — web, installable PWA, Android APK'],
    ['Asset coverage', 'Crypto, stocks (24 exchanges incl. DFM/Tadawul/NSE), ETFs, commodities, CME futures, indices, FX, US options, Web3 on-chain'],
    ['Proof', '29 consecutive days of public, hash-chained accuracy snapshots; 19,550 graded projections; no back-filling'],
    ['Users', 'None — pre-launch. Acquisition not started'],
    ['Revenue', 'US$0. Free/Pro/Ultimate plans and Stripe billing built, one switch from live'],
    ['Team', 'Solo founder. First hires are the primary use of this round'],
    ['Capital raised to date', 'US$0 — bootstrapped'],
])

# ---- 2 problem ----
H1('2. Problem')
P('A professional markets terminal costs US$25,000 or more per year. That prices out the fastest-growing investor '
  'population in the world: first-generation retail investors across the GCC, India and Pakistan. What fills the gap is '
  'a market of fragmented apps and Telegram tipsters selling certainty that cannot exist — guaranteed calls, back-filled '
  'win rates, cherry-picked screenshots.')
B('No mainstream retail product publishes an auditable record of when it was wrong.')
B('Regulators across the GCC and India are actively tightening on fake-signal and unlicensed advisory apps.')
B('The information asymmetry between institutions and first-time retail investors is one of the quietest inequities in emerging-market finance.')

# ---- 3 solution ----
H1('3. Solution — the product, as it exists today')
H2('Coverage')
B('Crypto — tick-by-tick live streaming, 24/7')
B('Equities — 24 exchanges including DFM, Tadawul, NSE, BSE, JSE, LSE, NYSE/Nasdaq')
B('ETFs, commodities, CME futures, world indices, FX, and US options chains')
B('Web3 — DeFi TVL by chain, protocol rankings, gas, market structure, full token universe')
B('Displayed in 24 local currencies; Arabic and right-to-left interface shipped')
H2('Intelligence')
B('Plain-language AI analysis of any asset, with an educational glossary layer for first-time investors')
B('Quantra Score — a probabilistic 1–99 read, weighted by what has actually worked out-of-sample for that asset')
B('Calibrated projection bands (50% and 80%) that self-correct against measured live outcomes')
B('Movers radar — modeled odds of a 2–40% move across the market, 1 hour to 30 days, with opt-in alerts')
B('Portfolio risk analytics (correlation matrix, diversification score), paper trading, watchlists, alerts, daily brief')
H2('Platform')
B('Web application, installable PWA, Android APK; accounts mandatory')
B('Developer API with per-workspace keys; team workspaces with shared watchlists')
B('40-check CI on every release; self-diagnostics every 12 minutes; public system status page')

# ---- 4 moat ----
H1('4. The Moat — verifiable honesty')
P('Every competitor claims accuracy. Quantra proves it, or proves the opposite, in public.')
B('Every projection is a probability with the downside shown — never a promise.')
B("Each day's projections are hash-chained (SHA-256) to the previous day's. Editing or back-dating any record breaks the chain publicly.")
B('As of today: 29 consecutive days, 19,550 graded projections, 15,800+ calibration observations — growing daily.')
B('The forecast bands self-calibrate: measured coverage feeds back into the engine so an "80% band" converges on a true 80%.')
P('A competitor can copy the feature surface in months. They cannot copy months of public, verified accuracy history — '
  'that only accrues in real time. As regional regulators close in on fake-signal apps, the transparency-first product '
  'becomes the compliant default rather than a marketing angle.', bold=True)

# ---- 5 market ----
H1('5. Market')
B('India & South Asia — NSE publicly reports over 100 million registered investors, compounding; the largest F&O market in the world by contract volume.')
B('GCC — record retail participation on Tadawul and DFM, driven by national digitisation programmes and a continuing IPO pipeline.')
B('Crypto & Web3 — MENA is among the fastest-growing adoption regions (Chainalysis); the UAE has purpose-built regimes (VARA, ADGM).')
B('Wedge — the underserved multi-asset retail analyst. Expansion — team plans, developer API, licensed-data premium tiers, and Africa (JSE already live).')
P('Note on sizing: we deliberately do not present a top-down TAM multiplication. The year-3 projection reaches a small '
  'fraction of one percent of the GCC + India retail investor base; growth in this plan is marketing-led, not '
  'penetration-capped.', size=9.5, italic=True, color=MUT)

# ---- 6 business model ----
H1('6. Business Model')
table([
    ['Plan', 'Price (US$/mo)', 'Includes'],
    ['Free', '0', 'Live boards, basic analysis, 25-item watchlist'],
    ['Pro', '9.99', 'AI verdicts (300/day), intraday, exports, 200-item watchlist'],
    ['Ultimate', '24.99', 'AI verdicts (1,500/day), developer API, team workspace, 1,000-item watchlist'],
])
B('Primary customer: the self-directed retail investor in the GCC and South Asia.')
B('Revenue: monthly subscription (freemium → Pro/Ultimate). All three plans and Stripe billing are already built; enforcement is one environment flag.')
B('Blended ARPU at the modelled 80/20 Pro:Ultimate mix ≈ US$13.00/month.')
B('Assumed conversion: 4% of active free users — conservative against the 2–6% prosumer range.')
B('Assumed lifetime: at 5% monthly churn, ~20 months average → LTV ≈ US$260 per paying customer (gross).')
B('Future streams (modelled as upside, not relied upon): developer API tier, team/enterprise seats, licensed-data premium.')

# ---- 7 GTM ----
H1('7. Go-to-Market')
B('Built-in loop: shareable analysis snapshots — any user can publish a read as a public link that unfurls a branded preview and deep-links recipients into the product. Live today.')
B('Language: Arabic/RTL shipped — most honest competitors in this space are English-only.')
B('Channels: performance marketing in the GCC and India, finance-creator partnerships, app-store presence, and the public track record as the credibility asset in every campaign.')
B('Sequence: launch to a small cohort → prove retention and conversion → scale spend against a measured CAC.')
B('The track record page is the marketing: it is the one claim in this category a competitor cannot copy or fake.')

# ---- 8 competition ----
H1('8. Competition')
table([
    ['Category', 'Examples', 'Where Quantra differs'],
    ['Institutional terminals', 'Bloomberg, Refinitiv', 'US$25k+/yr, institutional-only. Quantra is retail-priced and retail-designed.'],
    ['Charting platforms', 'TradingView', 'Excellent charts, but analysis is user-generated and ungraded. Quantra generates the read and grades itself publicly.'],
    ['Retail brokers', 'Zerodha, eToro, local GCC brokers', 'Execution-first, analysis thin. Quantra is analysis-first and broker-neutral (bring your own broker).'],
    ['Signal apps / tipsters', 'Telegram groups, "AI signal" apps', 'Sell certainty, publish no auditable record, and are the target of tightening regulation. Quantra is the structural opposite.'],
])
P('Our defensible position is not features — it is the accumulating public accuracy record plus the regulatory '
  'direction of travel in our home markets.', bold=True)

# ---- 9 team ----
H1('9. Team')
P('Eshan Thanvi — Founder. Built Quantra AI end-to-end AI-assisted: eight asset classes, an analysis engine, '
  'personalization, teams, billing, PWA and Android, with CI and self-diagnostics — the feature surface of a Series-A '
  'company, at seed cost. This is the capital efficiency the current AI cycle enables, demonstrated rather than claimed.')
P('Stated plainly: Quantra is a solo founder today, with no co-founder. That is the single biggest gap in this plan.', bold=True, color=RED)
H2('Hiring plan (funded by this round)')
table([
    ['Role', 'When', 'Why it is key'],
    ['Co-founder / founding engineer', 'Immediate', 'Removes key-person risk; the primary gap'],
    ['2 × engineers', 'Months 1–6', 'Data adapters, mobile, scale'],
    ['Market-data specialist', 'Months 1–3', 'Licensed exchange feeds (Gulf, NSE/BSE F&O)'],
    ['Growth lead', 'Months 1–3', 'GCC + India go-to-market'],
    ['Compliance counsel (fractional)', 'Months 1–6', 'ADGM/UAE footing for a paid research product'],
])

# ---- 10 financials ----
H1('10. Financials')
P('Full model: Quantra_AI_Financial_Model.xlsx — budget-driven with live formulas and a Bear/Base/Bull switch; '
  'change any driver and the model recalculates. Base case below.', size=9.5, italic=True, color=MUT)
table([
    ['Metric (Base case)', 'Year 1', 'Year 2', 'Year 3'],
    ['Total Customers (paying, end of year)', '~990', '~2,520', '~4,120'],
    ['Total Revenue (US$)', '~92,000', '~296,000', '~541,000'],
    ['Total Expense (US$)', '~564,000', '~1,197,000', '~1,804,000'],
    ['EBITDA (US$)', '~(473,000)', '~(901,000)', '~(1,263,000)'],
    ['ARR run-rate (end of year, US$)', '~155,000', '~393,000', '~642,000'],
    ['Cash balance (end of year, US$)', '~2,027,000', '~1,126,000', '~(137,000)'],
])
H2('Scenarios')
table([
    ['Scenario', 'Y3 Revenue', 'Y3 ARR', 'LTV:CAC (Y1)', 'What it assumes'],
    ['Bear', '~95,000', '~109,000', '0.5x', 'CAC $9, conversion 2.5%, churn 7%, weak share loop'],
    ['Base', '~541,000', '~642,000', '2.3x', 'CAC $6, conversion 4%, churn 5%, organic 35%'],
    ['Bull', '~2,707,000', '~3,284,000', '10.5x', 'CAC $4.50, conversion 6%, churn 3.5%, organic 60%, higher pricing'],
])
H2('Unit economics — stated honestly')
P('The Base case runs at roughly 2.3x LTV:CAC in year 1, drifting to ~1.5x by year 3 as CAC inflates. That is below '
  'the 3:1 bar investors look for, and we have not tuned the assumptions until 3:1 appeared.', bold=True, color=RED)
P('The reason is structural, not evasive: CAC and conversion are the two numbers a pre-launch company cannot know — '
  'we have no users. So the first milestone of this round is to MEASURE both on a small cohort (~US$25k of spend) '
  'before scaling anything. Marketing scales only once the ratio clears 3:1. If it cannot be made to clear, the honest '
  'outcome is a smaller, slower company — not a larger burn.')
P('The levers, in order of impact: conversion (4% → 5% takes Base to ~2.8x); organic share from the in-product share '
  'loop (35% → 60%); CAC mix (Gulf is expensive, India is cheap); and churn, which drives LTV directly. The Bull case '
  'is all four landing well (10.5x); the Bear case is them missing (0.4x). Both are in the model, not hidden.')
H2('Key and critical assumptions')
B('Budget-driven: marketing spend is an input; signups fall out of it at the stated CAC. Growth is bought at a stated price, never assumed.')
B('Free → paid conversion: 4% Base (prosumer range 2–6%).')
B('Pricing: Pro US$9.99 / Ultimate US$24.99 — both already built in the product. Blended ARPU ≈ US$13.00.')
B('Churn: 5%/month paying (LTV ≈ US$252 net of fees), 8%/month free.')
B('Blended paid CAC per signup: US$6 (Y1) inflating to US$9 (Y3); organic adds 35% on top at zero CAC.')
B('Costs step in three tiers across Y1/Y2/Y3: team, licensed market data, infrastructure, marketing, G&A.')
B('No revenue is modelled before the round closes; there is no historical revenue in the model.')
H2('Runway and the Series A — stated plainly')
P('The US$2.5M seed funds the plan into year 3; Base-case closing cash in year 3 is slightly negative (~US$137k). '
  'This plan assumes a Series A in year 3, raised against a ~US$640k ARR run-rate, a live product and a multi-year '
  'public accuracy record. If no Series A is raised, marketing and hiring throttle back to hold cash positive — growth '
  'slows and the company survives. EBITDA is negative across all three years in Base; breakeven is not claimed inside '
  'the seed window. In the Bull case the company reaches EBITDA-positive in year 2 — that is upside, not the plan.', color=RED)

# ---- 11 use of funds ----
H1('11. Use of Funds — US$2.5M')
table([
    ['Allocation', 'Share', 'US$', 'Detail'],
    ['Team', '35%', '875,000', '4 engineers + data specialist + growth lead; 24-month runway'],
    ['Data & infrastructure', '30%', '750,000', 'Licensed Gulf + NSE/BSE feeds, premium APIs, production hosting'],
    ['Growth & go-to-market', '20%', '500,000', 'GCC + India marketing, app-store launches, referral loops'],
    ['Regulatory & legal', '10%', '250,000', 'ADGM/UAE footing for a paid research product'],
    ['Reserve', '5%', '125,000', 'Contingency, audit, insurance'],
])

# ---- 12 risks ----
H1('12. Risks and Mitigations')
table([
    ['Risk', 'Mitigation'],
    ['No traction yet — the core unknown', 'Product is complete, so capital goes to acquisition, not building. Launch to a small cohort and prove retention before scaling spend.'],
    ['Key-person risk — solo founder', 'Co-founder/founding-engineer search is the first hire funded by this round. Codebase is documented, tested and CI-guarded.'],
    ['Data licensing cost and dependency', 'Feed adapters are written and provider-agnostic with graceful fallback; a provider change is a config change, not a rebuild.'],
    ['Regulatory — research/advisory licensing', '10% of the round is allocated to ADGM/UAE footing. The product is architected as probabilistic information, never advice, with public grading.'],
    ['Incumbent response (TradingView, brokers)', 'They can copy features; they cannot retroactively produce a public, hash-chained accuracy history — and their business models discourage publishing it.'],
    ['Model accuracy disappoints', 'Accuracy is published either way. Calibration is self-correcting; honesty is the product, so a mediocre month is disclosure, not an existential event.'],
])

# ---- 13 roadmap ----
H1('13. Roadmap')
table([
    ['Window', 'Milestones'],
    ['Months 0–3', 'Licensed Gulf + NSE/BSE feeds live (adapters already written); ADGM process starts; billing switched on; co-founder + first engineers hired.'],
    ['Months 3–9', 'GCC + India go-to-market; Play Store and App Store releases; referral and share loops scaled; first paid cohorts.'],
    ['Months 9–18', 'Backtesting playground; deeper portfolio analytics; developer API tier commercialised; Africa market adapters.'],
    ['Months 18–24', 'Revenue traction plus a multi-year public accuracy record — the defensible data story for the Series A.'],
])

d.add_paragraph()
P('This document contains forward-looking projections based on the stated assumptions. They are not forecasts or '
  'guarantees. Quantra AI has no revenue and no active users as of the date of this plan. Quantra provides '
  'probabilistic analytics and never investment advice.', size=8.5, italic=True, color=RED)

d.save(OUT)
print('SAVED:', OUT)
