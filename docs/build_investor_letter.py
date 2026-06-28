# -*- coding: utf-8 -*-
"""Investor cover letter for Quantra AI (by Eshan Thanvi) -> docs/Quantra_AI_Investor_Cover_Letter.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x1B, 0x9E, 0x6B)
INK = RGBColor(0x1A, 0x1F, 0x28)
MUTE = RGBColor(0x55, 0x5F, 0x6B)
BLUE = RGBColor(0x2F, 0x5C, 0xC8)

doc = Document()
st = doc.styles['Normal']; st.font.name = 'Segoe UI'; st.font.size = Pt(10.5); st.font.color.rgb = INK
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.8); sec.left_margin = sec.right_margin = Inches(1.0)


def para(runs, after=8, before=0, size=10.5, line=1.3, align=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = line
    if align is not None: p.alignment = align
    for (t, b, c, sz) in (runs if isinstance(runs, list) else [(runs, False, INK, size)]):
        r = p.add_run(t); r.font.bold = b; r.font.color.rgb = c; r.font.size = Pt(sz); r.font.name = 'Segoe UI'
    return p


def band():
    b = doc.add_paragraph(); b.paragraph_format.space_after = Pt(8)
    pb = b._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '8'); bt.set(qn('w:color'), '1B9E6B'); bt.set(qn('w:space'), '1')
    bd.append(bt); pb.append(bd)


# Letterhead
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
r = p.add_run('Quantra'); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = INK; r.font.name = 'Segoe UI'
r2 = p.add_run('AI'); r2.font.size = Pt(22); r2.font.bold = True; r2.font.color.rgb = GREEN; r2.font.name = 'Segoe UI'
para([('Eshan Thanvi  ·  Founder', False, MUTE, 9.5)], after=1)
para([('eshanthanvi@gmail.com  ·  quantra-ai.onrender.com', False, BLUE, 9.5)], after=4)
band()

para([('[Date]', False, MUTE, 10)], after=10)
para([('Dear [Investor Name],', False, INK, 10.5)], after=8)

para([('I am writing to introduce ', False, INK, 10.5), ('Quantra AI', True, INK, 10.5),
      (', a live market-analysis platform I founded and built, and to explore your interest in a strategic '
       'investment. Given your focus on [investor focus — e.g. fintech / early-stage technology / the GCC market], '
       'I believe Quantra is worth a closer look — and, unusually for an early-stage venture, it is already '
       'a working product you can evaluate today.', False, INK, 10.5)])

para([('Quantra AI is a ', False, INK, 10.5), ('live, professional-grade market-analysis platform', True, INK, 10.5),
      (' that puts the core of a $20,000-a-year trading terminal into any browser. It unifies real-time data across '
       '25 world exchanges — stocks, ETFs, indices, crypto, FX and commodities — with an advanced 16-study '
       'charting suite, an ', False, INK, 10.5), ('AI analyst', True, INK, 10.5),
      (' that explains any asset in plain language, and probabilistic forecasts ', False, INK, 10.5),
      ('independently calibrated to ~80% accuracy on live outcomes', True, INK, 10.5),
      (' (published transparently, not back-filled). On top of analysis sit the tools investors use daily: portfolio '
       'tracking, smart alerts, paper trading, a market calendar, and ', False, INK, 10.5),
      ('broker-linked execution', True, INK, 10.5),
      (' where users trade through their own regulated brokerage — Quantra never custodies funds.', False, INK, 10.5)])

para([('What sets this apart from most early-stage opportunities is that ', False, INK, 10.5),
      ('the execution risk is already removed', True, INK, 10.5),
      ('. The platform is built end-to-end, deployed in production, and open for you to inspect and validate at '
       'quantra-ai.onrender.com before any commitment. It was designed and shipped by a single founder — '
       'capital simply accelerates what has already been proven.', False, INK, 10.5)])

para([('I am raising ', False, INK, 10.5), ('AED 800,000', True, GREEN, 10.5),
      (' to fund the next phase: premium real-time data licensing, broker integrations, regulatory groundwork, and '
       'go-to-market. For a strategic partner, I am open to structuring this as an ', False, INK, 10.5),
      ('equity stake plus an exclusive / white-label arrangement', True, INK, 10.5),
      (' — designed around your goals, with a board or advisory seat if desired.', False, INK, 10.5)])

para([('If this is of interest, I would value 20–30 minutes to walk you through a live demonstration and the detailed '
       'proposal. I am confident a quick look at the working product will speak louder than any slide, and I would be '
       'glad to meet at your convenience, in person or by call. I will follow up shortly, but please feel free to '
       'reach me directly at the details below.', False, INK, 10.5)])

para([('Warm regards,', False, INK, 10.5)], after=14, before=4)
para([('Eshan Thanvi', True, INK, 11.5)], after=0)
para([('Founder, Quantra AI', False, MUTE, 9.5)], after=1)
para([('eshanthanvi@gmail.com  ·  quantra-ai.onrender.com', False, BLUE, 9.5)], after=10)

para([('Confidential — intended for the named recipient only. This is not an offer of securities and is subject '
       'to a definitive agreement.', False, MUTE, 8.5)], line=1.2)

doc.save(r"C:\Users\eshan\quantra-terminal\docs\Quantra_AI_Investor_Cover_Letter.docx")
print("OK investor cover letter")
