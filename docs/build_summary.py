# -*- coding: utf-8 -*-
"""Builds docs/Quantra_AI_Summary.docx — a short (1-2 page) version summary."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x1B, 0x9E, 0x6B)
INK = RGBColor(0x1A, 0x1F, 0x28)
MUTE = RGBColor(0x5C, 0x66, 0x72)
BLUE = RGBColor(0x2F, 0x5C, 0xC8)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Segoe UI'; st.font.size = Pt(10); st.font.color.rgb = INK
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.6)
sec.left_margin = sec.right_margin = Inches(0.8)


def para(text='', size=10, color=INK, bold=False, italic=False, after=4, before=0):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color; r.font.bold = bold; r.font.italic = italic; r.font.name = 'Segoe UI'
    return p


def head(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.size = Pt(11.5); r.font.bold = True; r.font.color.rgb = GREEN; r.font.name = 'Segoe UI'


def line(label, body):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(2)
    r = p.add_run(label + '  '); r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = INK
    r2 = p.add_run(body); r2.font.size = Pt(10); r2.font.color.rgb = MUTE


def bullet(label, body):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label); r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INK
    r2 = p.add_run('  ' + body); r2.font.size = Pt(9.5); r2.font.color.rgb = MUTE


# ---- header ----
para('Quantra AI', 22, INK, bold=True, after=0)
para('Version Summary — what each phase added', 12, GREEN, bold=True, after=4)
para('Live: quantra-ai.onrender.com   ·   Current build: assets v46 · service worker quantra-v28 · 2026',
     9, MUTE, after=2)
b = doc.add_paragraph(); b.paragraph_format.space_after = Pt(2)
pb = b._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '8'); bt.set(qn('w:color'), '1B9E6B'); bt.set(qn('w:space'), '1')
bd.append(bt); pb.append(bd)

head('Phase 1 · Foundation (v1.x)')
para('First deploy → responsive mobile → live data (crypto WebSocket + Finnhub) → candlesticks + dated '
     'Monte-Carlo projections → tick-by-tick seconds chart → RSI/MACD/Volume panes + Heikin Ashi → price '
     'alerts + drawing tools (trendline/level/Fibonacci) → super-admin oversight, audit log & footfall analytics.',
     9.5, MUTE, after=2)

head('Phase 2 · The 10-feature SaaS roadmap (v2.x)')
para('Portfolio tracker · Ask-Quantra AI analyst · server-side alerts + email · market discovery '
     '(heatmap/movers/breadth) · earnings/IPO/economic calendar · PWA + push · AI daily brief · Bollinger '
     '+ bar replay + saved layouts · paper trading · community ideas + leaderboard.', 9.5, MUTE, after=2)

head('Phase 3 · Data & forecast accuracy (v3.x)')
para('Premium feeds (FMP + marketaux) · forecasts recalibrated to 80% band coverage + live calibration '
     'widget · fixed the MSFT wrong-% bug · keep-warm (no cold starts) · 25 world exchanges + 32 indices + '
     'regional ETFs (currency follows exchange) · data-freshness line · Twelve Data + absolute up/down amount.',
     9.5, MUTE, after=2)

head('Phase 4 · Technical indicators (v4.x)')
para('McGinley / EMA / Parabolic SAR + Stochastic → full suite: Ichimoku, Supertrend, VWAP, Keltner, '
     'Donchian, ADX, Williams %R, CCI on both line & candlestick charts → enlarge mode + dynamic legend.',
     9.5, MUTE, after=2)

head('Phase 5 · Market status (v5.x)')
para('Open/closed badge → per-row dots → live-flipping status + open/close countdown → real-time US feed → '
     'timezone-accurate status → per-exchange holiday calendar → faster + honest freshness labels.',
     9.5, MUTE, after=2)

head('Latest work — build v40 → v46')
bullet('v40', 'Fixed mobile chart scrubbing — date/price tooltip now tracks on touch.')
bullet('Docs', 'Pitch deck (.pptx/.pdf) + full changelog (.docx/.pdf) generated.')
bullet('Accounts', 'Added then removed the signup approval gate; kept super-admin Delete.')
bullet('v44', 'Every signed-in user gets the Ultimate plan (sign-in still required).')
bullet('Repositioning', 'Removed all "educational/illustrative" copy; projections shown as 80%-calibrated.')
bullet('Track record', 'Short horizons (1/2/3-session) so past projections advance daily instead of stuck at zero.')
bullet('Email', 'Diagnostics + "Send test email" button to surface the real Resend delivery error.')
bullet('v45', 'Bring-your-own-broker trading (Alpaca) — paper-mode first, every order user-confirmed, funds never held by Quantra.')
bullet('v46', 'Fit-to-screen fix — no horizontal overflow on desktop or mobile.')
bullet('Admin', 'System-status rows for Twelve Data / FMP / marketaux keys.')

para('', 6)
para('Data sources: CoinGecko/CoinPaprika · Coinbase · Yahoo Finance · Finnhub · Twelve Data · FMP · '
     'marketaux · open.er-api — all paid sources key-gated with graceful fallback.', 8.5, MUTE, after=0)

doc.save(r"C:\Users\eshan\quantra-terminal\docs\Quantra_AI_Summary.docx")
print("OK summary docx")
